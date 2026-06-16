import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "academy-os_day23_code_review_and_scenario_tests.docx"
RESULTS = ROOT / "docs" / "scenario-test-results-day-23.json"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 17, "0F172A"),
        ("Heading 2", 13, "1F3A5F"),
        ("Heading 3", 11, "334155"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Academy OS Day 23 코드리뷰 및 사용자 시나리오 테스트")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0F172A")

    meta = doc.add_paragraph()
    meta.add_run("프로젝트: academy-os\n").bold = True
    meta.add_run("검토 기준: 데이터 원천 혼동, 권한 누락, 저장 실패 처리, 중복 상태값, 테스트 부족\n")
    meta.add_run("테스트 실행 방식: 현재 소스코드와 샘플 데이터를 대상으로 한 정적/데이터 검증")


def add_callout(doc, title, body, fill="F8FAFC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    p.add_run("\n" + body)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        set_cell_shading(hdr[idx], "E8EEF5")
        if widths:
            set_cell_width(hdr[idx], widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            color = None
            if str(value) == "FAIL":
                color = "9B1C1C"
            elif str(value) == "PASS":
                color = "166534"
            elif str(value) == "PENDING":
                color = "7A5A00"
            set_cell_text(cells[idx], value, color=color)
            if widths:
                set_cell_width(cells[idx], widths[idx])
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build():
    with RESULTS.open("r", encoding="utf-8") as f:
        results_payload = json.load(f)

    doc = Document()
    style_doc(doc)
    add_title(doc)

    summary = results_payload["summary"]
    add_callout(
        doc,
        "테스트 실행 결과",
        f"총 20개 시나리오 중 통과 {summary['pass']}개, 실패 {summary['fail']}개, 보류 {summary['pending']}개입니다. "
        "브라우저 자동화는 playwright-core 의존성 누락으로 실행하지 못했고, 현재 소스코드와 샘플 데이터를 대상으로 직접 검증했습니다.",
        fill="F4F6F9",
    )

    doc.add_heading("1. 우선순위별 코드리뷰 결과", level=1)
    findings = [
        ["P0", "학생 화면에서 모든 학생 선택 가능", "학생/학부모 개인정보 노출 위험", "학생 화면은 currentStudentId만 받도록 분리"],
        ["P0", "역할/권한 게이트 부재", "학생/학부모가 강사 기능에 접근할 수 있음", "instructor_owner, student, parent 권한 매트릭스 추가"],
        ["P1", "데이터 원천 혼동", "sampleData, localStorage, 화면 상태, 스냅샷이 섞임", "엔티티별 단일 원천과 마이그레이션 규칙 정의"],
        ["P1", "저장 실패 처리 불일치", "일부 저장 실패가 사용자에게 보이지 않음", "useStoredState에 try/catch와 전역 저장 오류 배너 추가"],
        ["P1", "숙제 상태값 중복", "status, studentStatus, teacherStatus가 충돌 가능", "display status는 파생값으로 계산"],
        ["P2", "후속조치 상태 enum 불일치", "문서와 코드의 상태값이 다름", "draft, scheduled, completed, canceled 중 하나로 통일"],
        ["P2", "모의 발송 오해 가능", "mock send가 실제 발송처럼 보일 수 있음", "발송 모의 기록으로 명확히 표기"],
        ["P2", "후속조치 중복 생성 가능", "같은 숙제로 보충 과제가 여러 개 생김", "taskType + sourceId + studentId 유니크 가드"],
        ["P3", "한글 인코딩 출력 불안정", "리뷰/수정 시 문구 깨짐 위험", "UTF-8 고정 및 UI 라벨 상수화"],
    ]
    add_table(doc, ["우선순위", "발견사항", "위험", "권장 수정"], findings, [900, 2200, 3000, 3260])

    doc.add_heading("2. 가장 먼저 고칠 항목", level=1)
    add_bullets(
        doc,
        [
            "학생 포털의 학생 선택 드롭다운을 제거하고, 강사용 미리보기 화면과 실제 학생 화면을 분리합니다.",
            "역할 기반 권한 게이트를 추가해 학부모는 읽기 전용, 학생은 본인 숙제/설문만 가능하도록 막습니다.",
            "localStorage 저장 실패를 중앙에서 감지하고, 화면에 저장 실패 상태를 표시합니다.",
            "숙제 상태값을 studentStatus와 teacherStatus 중심으로 정리하고, status는 파생값으로 계산합니다.",
            "후속조치 생성 시 같은 sourceId로 중복 생성되지 않도록 막습니다.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("3. 사용자 시나리오 테스트 20개 실행 결과", level=1)
    scenario_rows = [
        [item["id"], item["title"], item["status"], item["detail"]]
        for item in results_payload["results"]
    ]
    add_table(doc, ["ID", "시나리오", "결과", "상세"], scenario_rows, [650, 3000, 950, 4760])

    doc.add_heading("4. 실패/보류 항목 해석", level=1)
    add_bullets(
        doc,
        [
            "S01 실패: 학생 화면에서 전체 학생을 선택할 수 있어 실제 학생 로그인 화면으로 쓰기 어렵습니다.",
            "S02 보류: 학부모 전용 포털과 parent role이 아직 구현되지 않아 읽기 전용 검증을 완료할 수 없습니다.",
            "S12 실패: 같은 숙제로 후속조치를 여러 번 생성할 수 있습니다.",
            "S20 실패: localStorage 공통 저장 실패가 사용자에게 표시되지 않습니다.",
        ],
    )

    doc.add_heading("5. 다음 작업 제안", level=1)
    add_bullets(
        doc,
        [
            "Day 24-1: RoleGate 도입 및 사이드바 권한 필터링",
            "Day 24-2: 학생 화면을 실제 학생용과 강사용 미리보기로 분리",
            "Day 24-3: useStoredState 저장 실패 처리와 전역 저장 상태 배너 추가",
            "Day 24-4: Homework status 정규화 및 후속조치 중복 생성 방지",
            "Day 24-5: 위 20개 시나리오를 자동 테스트로 전환",
        ],
    )

    doc.core_properties.title = "Academy OS Day 23 코드리뷰 및 사용자 시나리오 테스트"
    doc.core_properties.subject = "academy-os MVP review"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
